import streamlit as st
import requests
from docx import Document
from datetime import datetime
import speech_recognition as sr
import base64
import json
import time 
from langdetect import detect
from deep_translator import GoogleTranslator

# Ollama API endpoint
OLLAMA_URL = "http://127.0.0.1:11434/api/generate"

# Sidebar with Chat History
st.sidebar.title("🏛️ Lawmate")
st.sidebar.write("Ask legal questions and generate legal documents.")

# Personalized Greeting Based on Time of Day
hour = datetime.now().hour
greeting = "Good morning! ☀️" if hour < 12 else "Good afternoon! 🌤️ " if hour < 18 else "Good evening! 🌙"
st.sidebar.subheader(greeting)

# Chat history storage
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "user_input" not in st.session_state:
    st.session_state.user_input = ""

# Display Chat History in Sidebar with Icons
st.sidebar.subheader("Chat History")
for message in reversed(st.session_state.chat_history[-10:]):
    if "*You:*" in message:
        st.sidebar.markdown(f"🧑‍⚖ *{message}*")
    else:
        st.sidebar.markdown(f"🤖 *{message}*")

# Function to set background from an uploaded file

def set_background():
    image_path = "D:\projects\chatbot\legal2.jpg"  # Ensure this file is in your project folder

    with open(image_path, "rb") as f:
        encoded_string = base64.b64encode(f.read()).decode()

    st.markdown(
        f"""
        <style>
            /* Set Full Background */
            .stApp {{
                background-image: url("data:image/jpeg;base64,{encoded_string}");
                background-size: cover;
                background-position: center;
                background-attachment: fixed;
            }}

            /* Make Sidebar Transparent */
            [data-testid="stSidebar"] {{
                background: rgba(0, 0, 0, 0.3) !important;  /* Adjust transparency */
                backdrop-filter: blur(10px);  /* Optional: Adds a blur effect */
            }}

            /* Set Input Fields and Buttons to Match Theme */
            .stTextArea, .stSelectbox, .stButton > button {{
                background: rgba(0, 0, 0, 0.6) !important;
                color: white !important;
                border-radius: 8px;
            }}

            /* Improve Text Visibility */
            h1, h2, h3, h4, h5, h6, p, label, span {{
                color: white !important;
            }}
        </style>
        """,
        unsafe_allow_html=True,
    )

# Apply background
set_background()

# Main Chatbot UI
st.title("🏛️ Lawmate")
st.write("Ask me any legal question!")

# Function to Update Input
def update_input():
    st.session_state.user_input = st.session_state.input_text

# Speech-to-Text Function with Loading Indicator
def recognize_speech():
    recognizer = sr.Recognizer()
    with sr.Microphone() as source:
        with st.spinner("🎙 Listening... Speak now."):
            recognizer.adjust_for_ambient_noise(source)
            audio = recognizer.listen(source)
    try:
        text = recognizer.recognize_google(audio)
        st.success(f"✅ Recognized: {text}")
        st.session_state.user_input = text
    except sr.UnknownValueError:
        st.error("❌ Could not understand audio")
    except sr.RequestError:
        st.error("⚠ Speech recognition service error")

# Voice Input Button
if st.button("🎤 Speak Your Question"):
    recognize_speech()

# Quick Action Buttons for Common Legal Topics
col1, col2, col3 ,col4= st.columns(4)
with col1:
    if st.button("📖 Contract Law"):
        st.session_state.user_input = "Tell me about contract law."
with col2:
    if st.button("𓍝 Civil Rights"):
        st.session_state.user_input = "Explain civil rights."
with col3:
    if st.button("🏠 Property Law"):
        st.session_state.user_input = "What are the basics of property law?"
with col4:
    if st.button("📜 Bharatiya Nyaya Sanhita"):
        st.session_state.user_input ="give equivalent Bharatiya Nyaya Sanhita number to ipc section "
        
# Text Area for User Input
st.text_area("Your Question:", key="input_text", height=100, value=st.session_state.user_input, on_change=update_input)

# Document Template Selection
doc_type = st.selectbox("Choose Legal Document Type (if needed)", ["None", "Contract", "Agreement", "Legal Notice", "Letter"])

#language changing
def translate_response(question, response):
    question_lang = detect(question)
    if question_lang == "en":
        return response
    translated_response = GoogleTranslator(source="auto", target=question_lang).translate(response)
    return translated_response


# AI Response Button
if st.button("Ask"):
    user_input = st.session_state.user_input
    if user_input:
        try:
            prompt_text = (
                #f" Provide a brief response and generate a well-structured legal document format with complete terms and conditions relevant to the selected document type. "
                #f"if any ipc section is mentioned also mention its equivalent Bharatiya Nyaya Sanhita number along with the ipc section number at the very beginning."
                f"Here is the question:\n\n{user_input}"
                f"explain in less than 500 words"
            )

            # Send Request to Ollama API

            payload = {
                "model": "lawmate",  # Ensure the model is installed
                "prompt": prompt_text,  # User input
                "stream": True  # Enable streaming response
            }

            # Make the API request with streaming
            try:
                response = requests.post(OLLAMA_URL, json=payload, stream=True)

                # Check if the response is valid
                if response.status_code == 200:
                    # Create a placeholder for the dynamic response
                    placeholder = st.empty()
                    legal_advice = ""

                    # Iterate over each line of the streamed response
                    for line in response.iter_lines():
                        if line:  # Ensure line is not empty
                            try:
                                # Decode each line to a string and load it as JSON
                                chunk = json.loads(line.decode("utf-8"))
                                
                                # Check if the 'response' field exists in the chunk
                                if "response" in chunk:
                                    # Add the new part of the response to the existing content
                                    legal_advice += chunk["response"]

                                    translated_advice = translate_response(user_input, legal_advice)
                                    
                                    # Update the placeholder with the new response part
                                    #placeholder.write(f"💡 **AI Response:** {legal_advice}")
                                    placeholder.write(f"💡 **AI Response:** {translated_advice}")
                                    
                                    # Optional: Introduce a small delay to make the output more readable
                                    time.sleep(0.1)
                                
                            except json.JSONDecodeError as e:
                                st.write(f"⚠ Error decoding JSON chunk: {e}")
                    
                else:
                    st.write(f"⚠ Error: Received status code {response.status_code}")
                    st.write(f"Response Text: {response.text}")

            except requests.exceptions.RequestException as e:
                st.write(f"⚠ Connection Error: {e}")


            # Store Chat History
            st.session_state.chat_history.append(f"*You:* {user_input}")
            st.session_state.chat_history.append(f"*Chatbot:* {legal_advice}")

            # Feedback Section
            st.subheader("💬 Feedback")
            st.write("How helpful was the response?")

            # Optional: Text Input for Detailed Feedback
            feedback_text = st.text_area("Any additional comments?", key="feedback_input")

            # Submit Feedback Button
            if st.button("Submit Feedback"):
                if st.session_state.feedback:
                    st.success("✅ Feedback submitted successfully!")
        
            # Store feedback data (this could be logged or saved in a database)
                feedback_data = {
                    "user_input": user_input,
                    "bot_response": legal_advice,
                    "rating": st.session_state.feedback,
                    "comments": feedback_text
                }

            # Save feedback (for now, print it, but this can be written to a file or database)
                print(json.dumps(feedback_data,indent=4))

            # Generate Formatted Document
            if doc_type != "None":
                doc = Document()
                doc.add_heading(f"{doc_type}", level=1)
                doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
                doc.add_paragraph(f"Subject: {user_input}")
                
                if doc_type == "Contract":
                    doc.add_paragraph("PARTIES:\n[Party 1 Name]\n[Party 2 Name]\n")
                    doc.add_paragraph("TERMS AND CONDITIONS:")
                    doc.add_paragraph(legal_advice)
                elif doc_type == "Agreement":
                    doc.add_paragraph("AGREEMENT BETWEEN:\n[Party 1] and [Party 2]")
                    doc.add_paragraph("TERMS:")
                    doc.add_paragraph(legal_advice)
                elif doc_type == "Legal Notice":
                    doc.add_paragraph("TO: [Recipient Name]")
                    doc.add_paragraph("SUBJECT: Legal Notice")
                    doc.add_paragraph(legal_advice)
                elif doc_type == "Letter":
                    doc.add_paragraph("[Your Name]\n[Your Address]\n[Date]\n")
                    doc.add_paragraph("Dear [Recipient],")
                    doc.add_paragraph(legal_advice)
                    doc.add_paragraph("Sincerely,\n[Your Name]")
                
                # Save and Provide Download Link
                doc_filename = f"{doc_type}Legal_Document{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
                doc.save(doc_filename)
                
                with open(doc_filename, "rb") as file:
                    st.download_button(
                        label="📄 Download Legal Document", 
                        data=file, 
                        file_name=doc_filename, 
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

        except Exception as e:
            st.write("❌ Error:", e)
    else:
        st.write("⚠ Please enter a question.")
